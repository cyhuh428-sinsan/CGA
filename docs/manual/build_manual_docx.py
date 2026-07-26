from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "dist"
FONT = "Malgun Gothic"
HEADER_LABEL = "CGA Studio 메뉴얼"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(95, 99, 104)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=10.5, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^]]+\]\([^)]+\))")


def add_inline_runs(paragraph, text, size=10.5, color=None, bold=False, italic=False):
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        part_bold = bold
        part_italic = italic
        display = part
        if part.startswith("**") and part.endswith("**"):
            display = part[2:-2]
            part_bold = True
        elif part.startswith("`") and part.endswith("`"):
            display = part[1:-1]
            part_bold = True
        else:
            link = re.fullmatch(r"\[([^]]+)\]\(([^)]+)\)", part)
            if link:
                display = link.group(1)
        run = paragraph.add_run(display)
        set_font(run, size=size, color=color, bold=part_bold, italic=part_italic)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 14, 7), ("Heading 3", 12, DARK_BLUE, 10, 5)):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run(HEADER_LABEL), size=8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("CGA Studio"), size=8.5, color=GRAY)


def add_text(doc, text, style=None, bold=False, italic=False, color=None):
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text, color=color, bold=bold, italic=italic)
    return p


def add_table(doc, rows):
    if not rows:
        return
    count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=count)
    table.style = "Table Grid"
    widths = [9360 // count] * count
    widths[-1] += 9360 - sum(widths)
    set_table_width(table, widths)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, value.strip(), size=9.2, bold=row_index == 0, color=DARK_BLUE if row_index == 0 else None)
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def convert_markdown(source, output):
    doc = Document()
    style_document(doc)
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    title_added = False
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[2:].strip())
            set_font(run, size=24, color=DARK_BLUE, bold=True)
            title_added = True
            index += 1
            continue
        heading = re.match(r"^(##{1,3})\s+(.*)$", line)
        if heading:
            add_text(doc, heading.group(2).strip(), style=f"Heading {len(heading.group(1)) - 1}", bold=True)
            index += 1
            continue
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            add_inline_runs(p, line[1:].strip(), size=9.5, color=GRAY, italic=True)
            index += 1
            continue
        image_match = re.match(r"^!\[([^]]*)\]\(([^)]+)\)$", line)
        if image_match:
            image_path = (source.parent / image_match.group(2)).resolve()
            if image_path.exists():
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(6.25))
                if image_match.group(1).strip():
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(8)
                    set_font(caption.add_run(image_match.group(1).strip()), size=9, color=GRAY, italic=True)
            else:
                add_text(doc, f"[이미지 파일 없음: {image_match.group(2)}]", color=GRAY, italic=True)
            index += 1
            continue
        if line.startswith("|"):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                current = lines[index].strip()
                values = [part.strip() for part in current.strip("|").split("|")]
                if not all(re.fullmatch(r"[-: ]+", value or " ") for value in values):
                    rows.append(values)
                index += 1
            add_table(doc, rows)
            continue
        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if list_match:
            is_numbered = list_match.group(2)[0].isdigit()
            p = doc.add_paragraph(style=None if is_numbered else "List Bullet")
            if is_numbered:
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(4)
            marker = f"{list_match.group(2)} " if is_numbered else ""
            add_inline_runs(p, marker + list_match.group(3).strip(), size=10.5)
            index += 1
            continue
        if re.match(r"^상태:", line) or re.match(r"^대상:", line):
            add_text(doc, line, bold=True, color=GRAY)
            index += 1
            continue
        add_text(doc, line)
        index += 1
    if not title_added:
        add_text(doc, source.stem, style="Heading 1", bold=True)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


MANUALS = {
    "cga-user-manual": {"ko": "CGA 사용자 설명서", "default": "CGA User Manual"},
    "cga-getting-started": {"ko": "CGA Getting Started", "default": "CGA Getting Started"},
    "cga-nlu-guide": {"ko": "CGA NLU 활용 가이드", "default": "CGA NLU Guide"},
}
LANGUAGES = ("ko", "en", "zh-CN", "ja", "vi", "fr", "de")
FONT_BY_LANGUAGE = {
    "ko": "Malgun Gothic", "en": "Arial", "zh-CN": "Microsoft YaHei",
    "ja": "Yu Gothic", "vi": "Arial", "fr": "Arial", "de": "Arial",
}
HEADER_BY_LANGUAGE = {
    "ko": "CGA Studio 메뉴얼", "en": "CGA Studio Manual", "zh-CN": "CGA Studio 手册",
    "ja": "CGA Studio マニュアル", "vi": "Hướng dẫn CGA Studio",
    "fr": "Manuel CGA Studio", "de": "CGA Studio Handbuch",
}


def main():
    global FONT, HEADER_LABEL
    OUT.mkdir(exist_ok=True)
    for manual_dir, names in MANUALS.items():
        for language in LANGUAGES:
            FONT = FONT_BY_LANGUAGE[language]
            HEADER_LABEL = HEADER_BY_LANGUAGE[language]
            source_name = "README.md" if language == "ko" else f"README.{language}.md"
            source = ROOT / manual_dir / source_name
            output_stem = names["ko"] if language == "ko" else f"{names['default']} ({language})"
            convert_markdown(source, OUT / f"{output_stem}.docx")

if __name__ == "__main__":
    main()
